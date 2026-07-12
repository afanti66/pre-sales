#!/usr/bin/env python3
"""DOCX方案书生成器 - 将各章节合并为一份完整的Word文档"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, re

NFS = "/home/afanti/projects/sales-agent-dashboard/nfs-data/projects/user-management"
OUTPUT = f"{NFS}/artifacts/solution/proposal_v1.docx"

doc = Document()

# 样式设置
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# 封面
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('某单位统一用户管理与身份认证系统\n技术方案')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('版本：V1.0\n日期：2026年7月')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# 目录页
doc.add_heading('目  录', level=1)
toc_items = [
    '第一章 项目概述',
    '第二章 需求分析',
    '  2.1 组织架构管理需求',
    '  2.2 用户生命周期管理需求',
    '  2.3 身份认证需求',
    '  2.4 性能需求对应',
    '第三章 总体方案设计',
    '  3.1 总体架构',
    '  3.2 部署方案',
    '  3.3 系统集成方案',
    '第四章 详细设计',
    '  4.1 用户管理服务',
    '  4.2 身份认证服务',
    '  4.3 权限管理服务',
    '  4.4 安全审计服务',
    '第五章 实施方案',
    '第六章 预算方案',
    '附录A 接口规范'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# 读取各章节markdown并转换
md_files = [
    ("第一章 项目概述", f"{NFS}/input/requirements.md"),
    ("第二章 需求分析", f"{NFS}/artifacts/design/001_requirement.md"),
    ("第三章 总体方案设计", f"{NFS}/artifacts/design/002_architecture.md"),
    ("第四章 详细设计", f"{NFS}/artifacts/design/003_subproject.md"),
    ("第六章 预算方案", f"{NFS}/artifacts/design/006_budget.md"),
]

for chapter_title, md_path in md_files:
    if not os.path.exists(md_path):
        doc.add_heading(chapter_title, level=1)
        doc.add_paragraph('（内容待补充）')
        doc.add_page_break()
        continue

    with open(md_path, 'r') as f:
        content = f.read()

    doc.add_heading(chapter_title, level=1)

    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_rows = []

    for line in lines:
        # 代码块
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph('```')
            else:
                in_code_block = True
                doc.add_paragraph('```')
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(9)
            continue

        # 标题
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # 表格
        elif '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if table_rows and line.replace('|', '').replace('-', '').replace(':', '').strip() == '':
                continue
            table_rows.append(cells)
        else:
            # 如果之前有表格行，先输出表格
            if table_rows and len(table_rows) > 1:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        table.cell(i, j).text = cell_text
                doc.add_paragraph()
                table_rows = []

            # 普通文本
            if line.strip():
                # 处理加粗 **text**
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                if text.startswith('- '):
                    p = doc.add_paragraph(text[2:], style='List Bullet')
                else:
                    doc.add_paragraph(text)

    # 处理未闭合的表格
    if table_rows and len(table_rows) > 1:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text

    doc.add_page_break()

doc.save(OUTPUT)
print(f"✅ DOCX方案书已生成: {OUTPUT}")
print(f"   文件大小: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
